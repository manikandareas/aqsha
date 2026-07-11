"""Ekspor hasil analisis ke file unduhan (fase 5): .docx (tabel + interpretasi Bab 4),
.xlsx (tabel mentah), .sav (dataset untuk SPSS).

Dipanggil dari sandbox via `codeRun`; payload = grup blok terkumpul (tabel/verdict/figur)
yang sudah dipersist di `analysis_result_blocks`. Tak butuh jaringan.
"""

from __future__ import annotations

import base64
import os
from typing import Any


def _fmt(value: Any) -> str:
    if value is None:
        return "-"
    if isinstance(value, float):
        if value != value:  # NaN
            return "-"
        if value.is_integer():
            return str(int(value))
        return f"{value:.3f}".replace(".", ",")
    return str(value)


def export_docx(payload: dict, out_path: str) -> str:
    """Dokumen Word: judul + per grup (tabel bergaya SPSS, figur, interpretasi)."""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(payload.get("title") or "Hasil Analisis Data", level=1)

    for group in payload.get("groups", []):
        doc.add_heading(group.get("title") or "Analisis", level=2)
        for tbl in group.get("tables", []):
            if tbl.get("title"):
                doc.add_paragraph(tbl["title"], style="Intense Quote")
            columns = tbl.get("columns", [])
            rows = tbl.get("rows", [])
            table = doc.add_table(rows=1, cols=max(1, len(columns)))
            table.style = "Light Grid Accent 1"
            for i, col in enumerate(columns):
                table.rows[0].cells[i].text = str(col)
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    if i < len(cells):
                        cells[i].text = _fmt(val)
            for note in tbl.get("notes", []):
                p = doc.add_paragraph(str(note))
                p.runs[0].italic = True if p.runs else False
        for fig in group.get("figures", []):
            png = fig.get("png")
            if not png:
                continue
            tmp = f"{out_path}.fig.png"
            with open(tmp, "wb") as fh:
                fh.write(base64.b64decode(png))
            doc.add_picture(tmp, width=Inches(5.5))
            os.remove(tmp)
            if fig.get("caption"):
                doc.add_paragraph(str(fig["caption"]))
        decisions = group.get("decisions", [])
        if decisions:
            doc.add_paragraph("Interpretasi:").runs[0].bold = True
            for d in decisions:
                text = d.get("interpretation") or d.get("label") or ""
                if text:
                    doc.add_paragraph(text, style="List Bullet")

    doc.save(out_path)
    return out_path


def export_xlsx(payload: dict, out_path: str) -> str:
    """Workbook Excel: satu sheet per tabel (nama unik ≤31 char)."""
    from openpyxl import Workbook

    wb = Workbook()
    wb.remove(wb.active)
    used: set[str] = set()
    idx = 0
    for group in payload.get("groups", []):
        for tbl in group.get("tables", []):
            idx += 1
            base = (tbl.get("title") or f"Tabel {idx}")[:28]
            name = base
            n = 1
            while name in used or not name:
                name = f"{base[:26]} {n}"
                n += 1
            used.add(name)
            ws = wb.create_sheet(title=name)
            ws.append([str(c) for c in tbl.get("columns", [])])
            for row in tbl.get("rows", []):
                ws.append([_fmt(v) for v in row])
    if not wb.sheetnames:
        wb.create_sheet(title="Kosong")
    wb.save(out_path)
    return out_path


def export_sav(data_path: str, out_path: str) -> str:
    """Dataset → .sav (SPSS). Value/variable labels ikut bila sumbernya .sav/.dta."""
    import pyreadstat

    from .io import load_dataset

    df = load_dataset(data_path)
    # Nama kolom SPSS: alfanumerik/underscore, awali huruf, ≤64 char (hindari error write_sav).
    rename = {}
    for col in df.columns:
        safe = "".join(ch if (ch.isalnum() or ch == "_") else "_" for ch in str(col))
        if safe and not safe[0].isalpha():
            safe = f"V_{safe}"
        rename[col] = (safe or "VAR")[:64]
    df = df.rename(columns=rename)
    pyreadstat.write_sav(df, out_path)
    return out_path


def export_files(payload: dict, formats: list[str], data_path: str | None, out_dir: str) -> dict:
    """Bangun file terminta → {format: path}. Format tak dikenal / sav tanpa data dilewati."""
    out: dict[str, str] = {}
    if "docx" in formats:
        out["docx"] = export_docx(payload, os.path.join(out_dir, "hasil-analisis.docx"))
    if "xlsx" in formats:
        out["xlsx"] = export_xlsx(payload, os.path.join(out_dir, "hasil-analisis.xlsx"))
    if "sav" in formats and data_path:
        out["sav"] = export_sav(data_path, os.path.join(out_dir, "dataset-olahan.sav"))
    return out
