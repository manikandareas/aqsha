"""Fase 5 — ekspor .docx/.xlsx/.sav dari payload hasil analisis + dataset."""

from __future__ import annotations

from aqsha_stats.export import export_files


PAYLOAD = {
    "title": "Hasil Analisis Data",
    "groups": [
        {
            "title": "Uji validitas X1",
            "tables": [
                {
                    "title": "Item-Total Statistics",
                    "columns": ["Item", "r hitung", "r tabel", "Keputusan"],
                    "rows": [["X1.1", 0.612, 0.361, "Valid"], ["X1.2", 0.245, 0.361, "Tidak Valid"]],
                    "notes": ["a. df = n - 2"],
                }
            ],
            "decisions": [{"label": "Validitas X1.1", "verdict": "lolos", "interpretation": "Item X1.1 valid."}],
            "figures": [],
        }
    ],
}


def test_export_docx_and_xlsx(tmp_path):
    out = export_files(PAYLOAD, ["docx", "xlsx"], None, str(tmp_path))
    assert set(out) == {"docx", "xlsx"}

    from docx import Document

    doc = Document(out["docx"])
    texts = [p.text for p in doc.paragraphs]
    assert any("Uji validitas X1" in t for t in texts)
    assert any("valid" in t.lower() for t in texts)
    assert len(doc.tables) >= 1
    assert doc.tables[0].rows[0].cells[0].text == "Item"

    from openpyxl import load_workbook

    wb = load_workbook(out["xlsx"])
    ws = wb[wb.sheetnames[0]]
    assert [c.value for c in ws[1]][:2] == ["Item", "r hitung"]


def test_export_sav_roundtrip(fixture_path, tmp_path):
    out = export_files(PAYLOAD, ["sav"], fixture_path, str(tmp_path))
    assert "sav" in out

    import pyreadstat

    df, _meta = pyreadstat.read_sav(out["sav"])
    assert "X1" in df.columns and "Y" in df.columns
    assert len(df) == 100


def test_export_sav_skipped_without_data(tmp_path):
    out = export_files(PAYLOAD, ["sav"], None, str(tmp_path))
    assert out == {}
